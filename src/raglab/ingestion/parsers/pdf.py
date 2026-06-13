"""PDF and DOCX parsers (require the ``parsers`` extra)."""

from __future__ import annotations

from pathlib import Path

from raglab.core.registry import register
from raglab.core.types import Document


@register("parser", "pdf")
class PDFParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".pdf"

    def parse(self, path: Path) -> list[Document]:
        try:
            from pypdf import PdfReader
        except ImportError as e:  # pragma: no cover
            raise ImportError(
                "PDF parsing needs the 'parsers' extra: pip install 'raglab[parsers]'"
            ) from e
        reader = PdfReader(str(path))
        docs: list[Document] = []
        for page_no, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(
                    Document(
                        text=text,
                        metadata={"source": str(path), "page": page_no, "title": path.stem},
                    )
                )
        return docs


@register("parser", "docx")
class DocxParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> list[Document]:
        try:
            import docx
        except ImportError as e:  # pragma: no cover
            raise ImportError(
                "DOCX parsing needs the 'parsers' extra: pip install 'raglab[parsers]'"
            ) from e
        document = docx.Document(str(path))
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return [Document(text=text, metadata={"source": str(path), "title": path.stem})]
